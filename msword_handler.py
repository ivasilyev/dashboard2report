
import os
import logging
import tempfile
from docx import Document
from docx.shared import Mm
from pandas import DataFrame
from utils import datetime_now
from image_handler import ImageHandler


class MSWordHandler:
    def __init__(
            self,
            time_from: int,
            time_to: int,
            title: str = "Document Title",
    ):
        self.time_from = time_from
        self.time_to = time_to
        self.title = title

        self.document = Document()
        self._image_counter = 1
        self._table_counter = 1

    @property
    def _text_width(self):
        section = self.document.sections[0]
        return Mm((section.page_width - section.left_margin - section.right_margin) / 36000)

    @staticmethod
    def _assign_table_cells(cells, idx, text, is_header: bool = False):
        run = cells[idx].paragraphs[0].add_run(str(text))
        if is_header:
            run.bold = True

    def header(self, *args, **kwargs):
        self.document.add_heading(*args, **kwargs)

    def paragraph(self, *args, **kwargs):
        self.document.add_paragraph(*args, **kwargs)

    def embed_df(self, df: DataFrame, title: str = "", description: str = ""):
        if df.shape[1] == 0:
            logging.critical("The table has no rows")
            return
        logging.debug(f"Add table of shape '{df.shape}'")
        if len(description) > 0:
            self.paragraph(description.format(self._table_counter), style="Normal")
        self.paragraph()
        if len(title) == 0:
            title = f"Таблица {self._table_counter} – {df.name}"
        self.paragraph(title, style="Caption")
        self._table_counter += 1
        rendering_table = self.document.add_table(rows=1, cols=df.shape[1], style="Table Grid")
        cells = rendering_table.rows[0].cells
        for idx, i in enumerate(df.columns):
            self._assign_table_cells(cells, idx, i, is_header=True)
        for values in df.fillna("").astype(str).values:
            cells = rendering_table.add_row().cells
            for idx, i in enumerate(values):
                self._assign_table_cells(cells, idx, i)
        self.paragraph()

    def embed_image(self, handler: ImageHandler, title: str = "", description: str = ""):
        if not os.path.isfile(handler.file):
            logging.critical("The image file does not exist")
            return
        logging.debug(f"Add image '{handler.title}'")
        if len(description) > 0:
            self.paragraph(description.format(self._image_counter), style="Normal")
        self.paragraph()
        self.document.add_picture(handler.file, width=self._text_width)
        if len(title) == 0:
            title = f"Рисунок {self._image_counter} – {handler.title}"
        self.paragraph(title, style="Caption")
        self._image_counter += 1
        self.paragraph()

    def save(self, output_dir: str):
        output_file = os.path.join(output_dir, f"build-{datetime_now()}.docx")
        self.document.save(output_file)
        logging.info(f"Saved document: '{output_file}'")

    def render(self, output_dir: str):
        # Document creation routines
        pass

    def run(self, output_dir: str = ""):
        if len(output_dir) == 0:
            output_dir = tempfile.TemporaryDirectory().name
        logging.debug("Started document creation")
        self.render(output_dir)
        logging.debug("Finished document creation")
        self.save(output_dir)
